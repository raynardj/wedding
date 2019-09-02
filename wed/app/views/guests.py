from flask_appbuilder import ModelView, expose
from flask import render_template
from flask_appbuilder.models.sqla.interface import SQLAInterface

from wed.app.models.guests import guestTypeModel, guestModel, dinnerModel
from flask import request
from flask_login import current_user
from flask import g, send_file

from docx import Document
from datetime import datetime
from io import BytesIO
from docx.oxml.ns import qn
import random

class dinnerView(ModelView):
    route_base = "/dinner"
    datamodel =  SQLAInterface(dinnerModel)
    show_title = "桌信息"
    add_title = "添加新桌"
    edit_title = "编辑桌信息"
    list_title = "桌列表"
    label_columns = {
        "round_btn":"这一桌人",
        "remark":"备注","sn":"桌号",
        "guests":"宾客","guest_number":"宾客数量"}

    add_columns = ["sn","guests", "remark"]
    edit_columns = ["sn","guests","remark"]
    list_columns = ["round_btn","sn","remark","guest_number"]

    @expose("/round/<dinner_id>/")
    def roundTable(self, dinner_id):
        dinner = self.datamodel.get(int(dinner_id))
        return self.render_template("round_table.html", dinner = dinner)

    @expose("/map/")
    def dinnerMap(self):
        from .. import db
        dinners = db.session.query(dinnerModel).order_by(dinnerModel.sn.asc()).all()
        standing = db.session.query(guestModel).filter(guestModel.dinner_id==None).all()
        rdhash = hex(random.randint(10000, 20000))
        return self.render_template("map.html", dinners = dinners, standing = standing, rdhash=rdhash)

    @expose("/getmap/<rdhash>/")
    def getMap(self,rdhash):
        from .. import db
        dinners = db.session.query(dinnerModel).order_by(dinnerModel.sn.asc()).all()
        f = BytesIO()
        doc = Document()
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        doc.add_heading("宾客座位名单", level = 0)
        for dinner in dinners:
            dh1, dh2, guests = self.dinnerText(dinner)
            doc.add_heading(*dh1)
            doc.add_heading(*dh2)
            doc.add_paragraph(guests)

        now_ = (int(datetime.now().timestamp()))
        doc.save(f)
        f.seek(0)
        return send_file(f,as_attachment=True, attachment_filename="wedding_map_%s.docx"%(now_))

    def dinnerText(self, dinner):
        dh1,dh2 = ("%s(%s)"%(dinner.sn, dinner.guest_number),3),(str(dinner.remark),4)
        guests = ",".join(list("%s %s"%(str(guest),"(%s)"%guest.remark if guest.remark else "") for guest in dinner.guests))
        return dh1, dh2, guests

class guestTypeView(ModelView):
    route_base = "/gtype"
    datamodel = SQLAInterface(guestTypeModel)
    show_title = "宾客类型"
    add_title = "添加宾客类型"
    edit_title = "编辑宾客类型"
    list_title = "宾客大类"
    label_columns = {"name":"类型名称","remark":"备注","guests":"包含宾客","gcount":"总人数"}

    add_columns = ["name","remark"]
    edit_columns = ["name", "remark"]
    list_columns = ["name","guests","gcount","remark"]

    @expose("/detail/<gtype_id>/")
    def detail(self,gtype_id):
        gtype = self.datamodel.get(int(gtype_id))
        login = True if hasattr(g.user, "id") else False
        return self.render_template("gtype_detail.html", gtype = gtype, login = login)

    @expose("/all/")
    def all(self):
        groups = self.datamodel.session.query(guestTypeModel).all()
        login = True if hasattr(g.user, "id") else False
        return self.render_template("gtype_all.html", groups = groups, login = login)


class guestView(ModelView):
    route_base = "/guest"
    datamodel = SQLAInterface(guestModel)
    show_title = "宾客"
    add_title = "添加宾客"
    edit_title = "修改宾客信息"
    list_title = "宾客编辑列表"
    label_columns = {"fullname":"宾客全名","type":"宾客种类","babies":"携带宝宝数","phone":"手机",
                     "ishotel":"需安排住宿","ispickup":"需车辆接送","hotel":"酒店详情","pickup":"接送详情",
                     "invitation":"已发请柬","remark":"备注", "redpack":"红包哟",
                     }
    add_columns = ["fullname", "type","plus","babies","phone","invitation",
                   "remark","ishotel","hotel","ispickup","pickup"]
    edit_columns = ["fullname", "type","plus", "babies", "phone","redpack", "invitation",
                   "remark", "ishotel", "hotel", "ispickup", "pickup"]
    list_columns = ["fullname", "type","remark",  "phone", "invitation",
                     "ishotel",  "ispickup","babies",]

    @expose("/detail/<guest_id>/")
    def detail(self,guest_id):
        guest = self.datamodel.get(int(guest_id))
        login = True if hasattr(g.user,"id") else False
        return self.render_template("guest_detail.html",guest = guest , login = login)

    @expose("/evite/<guest_id>/")
    def evite(self, guest_id):
        guest = self.datamodel.get(int(guest_id))
        login = True if hasattr(g.user, "id") else False
        return self.render_template("invitation.html", guest=guest, login=login)

    @expose("/all/")
    def all(self):
        guests = self.datamodel.session.query(guestModel).all()
        guest_total = self.datamodel.session.query(guestModel).count()+sum(list(g.plus for g in guests))
        return self.render_template("all.html",guests = guests, guest_total = guest_total)

    @expose("/babies/")
    def babies(self):
        babyguest = self.datamodel.session.query(guestModel).filter(guestModel.babies>0).all()
        babycount = sum(list(bg.babies for bg in babyguest))
        print(babyguest)
        return self.render_template("babies.html", babyguest = babyguest, babycount = babycount)

    @expose("/hotel/")
    def hotels(self):
        guests = self.datamodel.session.query(guestModel).filter(guestModel.ishotel).all()
        guest_total = len(guests)
        return self.render_template("hotel.html", guests=guests, guest_total = guest_total)

    @expose("/pickup/")
    def pickups(self):
        guests = self.datamodel.session.query(guestModel).filter(guestModel.ispickup).all()
        guest_total = len(guests)
        return self.render_template("pickup.html", guests=guests, guest_total=guest_total)